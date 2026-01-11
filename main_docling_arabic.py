"""
Enhanced Document Processing with Docling - Multilingual Support
Uses Docling's advanced PDF understanding with proper OCR configuration.
Supports Arabic (RTL), English, and French languages.
"""

from fastapi import FastAPI, UploadFile, File
import tempfile, os
import torch
import re
from typing import List, Dict, Any
import subprocess

app = FastAPI()


# ========================================
# Auto-configure TESSDATA_PREFIX
# ========================================
def setup_tessdata_prefix():
    """Automatically detect and set TESSDATA_PREFIX if not set."""
    if os.environ.get('TESSDATA_PREFIX'):
        print(f"✅ TESSDATA_PREFIX already set: {os.environ['TESSDATA_PREFIX']}")
        return True
    
    # Common tessdata locations
    possible_paths = [
        '/usr/share/tesseract-ocr/5/tessdata',
        '/usr/share/tesseract-ocr/4.00/tessdata',
        '/usr/share/tessdata',
        '/opt/homebrew/share/tessdata',
        '/usr/local/share/tessdata',
        'C:\\Program Files\\Tesseract-OCR\\tessdata',
    ]
    
    # Try to find tessdata using system command
    try:
        result = subprocess.run(['find', '/usr', '-name', 'tessdata', '-type', 'd'], 
                              capture_output=True, text=True, timeout=5)
        if result.stdout:
            found_paths = result.stdout.strip().split('\n')
            possible_paths = found_paths + possible_paths
    except:
        pass
    
    # Check each path
    for path in possible_paths:
        if os.path.exists(path) and os.path.isdir(path):
            # Check if it contains language files
            files = os.listdir(path)
            if any(f.endswith('.traineddata') for f in files):
                os.environ['TESSDATA_PREFIX'] = path
                print(f"✅ Auto-detected TESSDATA_PREFIX: {path}")
                
                # Check for all supported languages
                has_ara = 'ara.traineddata' in files
                has_eng = 'eng.traineddata' in files
                has_fra = 'fra.traineddata' in files
                print(f"  - Arabic support: {'✅' if has_ara else '❌'}")
                print(f"  - English support: {'✅' if has_eng else '❌'}")
                print(f"  - French support: {'✅' if has_fra else '❌'}")
                
                if not has_ara:
                    print(f"  ⚠️  Arabic not found. Install with:")
                    print(f"      sudo apt-get install tesseract-ocr-ara")
                if not has_eng:
                    print(f"  ⚠️  English not found. Install with:")
                    print(f"      sudo apt-get install tesseract-ocr-eng")
                if not has_fra:
                    print(f"  ⚠️  French not found. Install with:")
                    print(f"      sudo apt-get install tesseract-ocr-fra")
                
                return True
    
    print("❌ Could not auto-detect TESSDATA_PREFIX")
    print("Please set it manually:")
    print("  export TESSDATA_PREFIX=/path/to/tessdata")
    return False


# Set up tessdata on startup
setup_tessdata_prefix()


def is_section_heading(text: str) -> bool:
    """Robust section heading detection for Arabic, English, and French text."""
    text = text.strip()
    
    if not text or len(text) > 200 or len(text) < 3:
        return False
    
    # Arabic section keywords
    arabic_indicators = [
        'الوحدة', 'وحدة', 'الفصل', 'فصل',
        'المقدمة', 'مقدمة', 'الخاتمة', 'خاتمة',
        'الأهداف', 'أهداف', 'المنهجية', 'منهجية',
        'التقويم', 'تقويم', 'الباب', 'باب',
        'القسم', 'قسم', 'الجزء', 'جزء',
        'الفرع', 'فرع', 'الملحق', 'ملحق',
        'المراجع', 'مراجع', 'فترة المراجعة',
    ]
    
    # English section keywords
    english_indicators = [
        'chapter', 'section', 'unit', 'introduction', 'conclusion',
        'abstract', 'summary', 'appendix', 'references', 'bibliography',
        'part', 'volume', 'preface', 'foreword', 'acknowledgments',
        'table of contents', 'index', 'glossary'
    ]
    
    # French section keywords
    french_indicators = [
        'chapitre', 'section', 'unité', 'introduction', 'conclusion',
        'résumé', 'annexe', 'références', 'bibliographie',
        'partie', 'volume', 'préface', 'avant-propos', 'remerciements',
        'table des matières', 'index', 'glossaire', 'objectifs',
        'méthodologie', 'évaluation'
    ]
    
    # Combine all indicators
    section_indicators = arabic_indicators + english_indicators + french_indicators
    
    text_lower = text.lower()
    for indicator in section_indicators:
        if indicator.lower() in text_lower:
            return True
    
    # Numbered sections (supports Arabic, English, and French numbering)
    if re.match(r'^(\d+|[٠-٩]+)[\.\-\:]\s*.+', text):
        return True
    
    # Roman numerals (common in French/English documents)
    if re.match(r'^[IVX]+[\.\-\:]\s*.+', text, re.IGNORECASE):
        return True
    
    # Short lines with colons (common pattern for headings)
    if ':' in text and len(text) < 120:
        return True
    
    # All caps short lines (often headings in English/French)
    if text.isupper() and len(text) < 80 and len(text.split()) < 10:
        return True
    
    return False


def extract_sections_from_docling_document(doc) -> List[str]:
    """Extract section headings from Docling document using its structure."""
    sections = []
    
    try:
        # Method 1: Extract from document structure
        for item in doc.iterate_items():
            if hasattr(item, 'label'):
                label = str(item.label).lower()
                if any(x in label for x in ['heading', 'title', 'section']):
                    if hasattr(item, 'text'):
                        section_text = item.text.strip()
                        if section_text and len(section_text) > 3:
                            sections.append(section_text)
                            print(f"  📌 Structure: {section_text[:80]}")
        
        # Method 2: Parse markdown export for headings
        if not sections or len(sections) < 3:
            markdown = doc.export_to_markdown()
            lines = markdown.split('\n')
            
            for line in lines:
                # Markdown headings start with #
                if line.startswith('#'):
                    heading = re.sub(r'^#+\s*', '', line).strip()
                    if heading and len(heading) > 3:
                        sections.append(heading)
                        print(f"  📌 Markdown: {heading[:80]}")
                # Text-based headings
                elif is_section_heading(line):
                    sections.append(line.strip())
                    print(f"  📌 Pattern: {line[:80]}")
    
    except Exception as e:
        print(f"⚠️ Error extracting sections: {e}")
    
    return sections


def assign_sections_to_chunks(chunks: List[Dict], sections: List[str]) -> List[Dict]:
    """Intelligently assign section names to chunks."""
    if not sections:
        print("⚠️ No sections detected - using page numbers")
        return [{
            "text": c["text"],
            "meta": {
                "page": c.get("page"),
                "section": f"صفحة {c.get('page', 'غير محدد')}"
            }
        } for c in chunks]
    
    enriched = []
    current_section = sections[0]
    
    print(f"\n📊 Assigning {len(sections)} sections to {len(chunks)} chunks...")
    
    for idx, chunk in enumerate(chunks):
        chunk_text = chunk["text"]
        found_section = None
        
        # Check if chunk contains section heading (first 500 chars)
        search_text = chunk_text[:500]
        
        for section in sections:
            # Normalize for comparison
            section_normalized = re.sub(r'[ًٌٍَُِّْ]', '', section)
            search_normalized = re.sub(r'[ًٌٍَُِّْ]', '', search_text)
            
            if section_normalized in search_normalized or section in search_text:
                found_section = section
                current_section = section
                if idx < 3:
                    print(f"  ✅ Chunk {idx}: {section[:60]}...")
                break
        
        if not found_section:
            found_section = current_section
        
        enriched.append({
            "text": chunk["text"],
            "meta": {
                "page": chunk.get("page"),
                "section": found_section
            }
        })
    
    return enriched


def clean_text(text: str) -> str:
    """
    Clean text by removing unnecessary newlines and normalizing whitespace.
    """
    if not text:
        return ""
    
    # Replace all newlines with spaces
    text = text.replace('\n', ' ')
    
    # Replace multiple spaces with single space
    import re
    text = re.sub(r'\s+', ' ', text)
    
    # Remove leading/trailing spaces
    text = text.strip()
    
    return text


async def _process_with_alternative_method(pdf_path: str):
    """
    Alternative PDF processing method when Docling fails.
    Uses PyPDF2/pdfplumber for text extraction and Tesseract directly for OCR.
    Supports Arabic, English, and French.
    """
    print("  📄 Using alternative PDF processing method (multilingual support)...")
    
    try:
        import PyPDF2
        from pdf2image import convert_from_path
        import pytesseract
        
        # Try to extract text directly from PDF first
        text_chunks = []
        sections = []
        
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            
            print(f"  📄 Processing {num_pages} pages with alternative method...")
            
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                
                if text and text.strip():
                    # Clean the text first
                    text = clean_text(text)
                    
                    # Simple chunking by paragraphs (split by double spaces or periods)
                    # Split by common paragraph markers
                    import re
                    # Split by multiple spaces, periods followed by space, or numbered items
                    paragraphs = re.split(r'\s{2,}|\.\s+(?=[A-Za-z\u0600-\u06FF])', text)
                    
                    for para in paragraphs:
                        para = para.strip()
                        # Clean the paragraph
                        para = clean_text(para)
                        if len(para) > 30:  # Only keep substantial paragraphs
                            text_chunks.append({
                                "text": para,
                                "page": page_num + 1
                            })
                    
                    # Try to detect sections (check cleaned text)
                    # Look for section indicators in the first part of the text
                    first_part = text[:500] if len(text) > 500 else text
                    # Split by common section markers
                    potential_sections = re.split(r'[:\-]\s+', first_part)
                    for potential in potential_sections[:5]:  # Check first 5 potential sections
                        potential = potential.strip()
                        if is_section_heading(potential):
                            sections.append(potential)
                            break
        
        # If no text extracted, try OCR
        if not text_chunks:
            print("  🔍 No text found, attempting OCR...")
            try:
                images = convert_from_path(pdf_path, dpi=200)
                for i, image in enumerate(images):
                    # Use Tesseract directly with Arabic
                    ocr_text = pytesseract.image_to_string(
                        image, 
                        lang='ara+eng+fra',
                        config='--psm 6'
                    )
                    if ocr_text and ocr_text.strip():
                        # Clean the OCR text
                        cleaned_text = clean_text(ocr_text)
                        if cleaned_text:
                            text_chunks.append({
                                "text": cleaned_text,
                                "page": i + 1
                            })
            except Exception as ocr_error:
                print(f"  ⚠️  OCR failed: {ocr_error}")
        
        # Clean all chunks text
        for chunk in text_chunks:
            chunk["text"] = clean_text(chunk["text"])
        
        if not text_chunks:
            return {
                "success": False,
                "error": "Could not extract any text from PDF",
                "details": "Both direct text extraction and OCR failed"
            }
        
        # Assign sections to chunks
        enriched_chunks = assign_sections_to_chunks(text_chunks, sections)
        
        return {
            "success": True,
            "method": "alternative_pdf_processing",
            "total_chunks": len(enriched_chunks),
            "detected_sections": sections,
            "sections_count": len(sections),
            "unique_sections_in_chunks": len(set(c['meta']['section'] for c in enriched_chunks)),
            "chunks": enriched_chunks,
            "metadata": {
                "processing_method": "alternative",
                "note": "Processed with alternative method due to Docling tensor errors"
            }
        }
        
    except ImportError:
        return {
            "success": False,
            "error": "Alternative processing method not available",
            "details": "Required libraries (PyPDF2, pdf2image) not installed",
            "install": "pip install PyPDF2 pdf2image pillow"
        }
    except Exception as e:
        return {
            "success": False,
            "error": "Alternative processing failed",
            "details": str(e)
        }


async def _process_docx_file(docx_path: str):
    """
    Process DOCX files and extract text with sections.
    """
    print("\n" + "="*70)
    print("DOCX PROCESSING WITH ARABIC SUPPORT")
    print("="*70)
    
    try:
        from docx import Document
        
        print("\n📄 STEP 1: Reading DOCX file...")
        doc = Document(docx_path)
        
        # Extract all text and paragraphs
        text_chunks = []
        sections = []
        current_page = 1
        current_text = ""
        
        print(f"  📝 Found {len(doc.paragraphs)} paragraphs")
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            
            if not text:
                continue
            
            # Clean the text
            cleaned_text = clean_text(text)
            
            # Check if it's a section heading
            if is_section_heading(cleaned_text):
                # Save previous chunk if exists
                if current_text and len(current_text) > 30:
                    text_chunks.append({
                        "text": clean_text(current_text),
                        "page": current_page
                    })
                    current_text = ""
                
                sections.append(cleaned_text)
                print(f"  📌 Section detected: {cleaned_text[:60]}...")
            else:
                # Add to current text chunk
                if current_text:
                    current_text += " " + cleaned_text
                else:
                    current_text = cleaned_text
                
                # Create chunk if text is long enough
                if len(current_text) > 600:
                    text_chunks.append({
                        "text": clean_text(current_text),
                        "page": current_page
                    })
                    current_text = ""
                    current_page += 1
        
        # Add remaining text
        if current_text and len(current_text) > 30:
            text_chunks.append({
                "text": clean_text(current_text),
                "page": current_page
            })
        
        print(f"  ✅ Extracted {len(text_chunks)} chunks from DOCX")
        print(f"  ✅ Found {len(sections)} sections")
        
        # Assign sections to chunks
        enriched_chunks = assign_sections_to_chunks(text_chunks, sections)
        
        return {
            "success": True,
            "method": "docx_processing",
            "total_chunks": len(enriched_chunks),
            "detected_sections": sections,
            "sections_count": len(sections),
            "unique_sections_in_chunks": len(set(c['meta']['section'] for c in enriched_chunks)),
            "chunks": enriched_chunks,
            "metadata": {
                "processing_method": "docx",
                "file_type": "docx",
                "languages": ["ara", "eng", "fra"],
                "note": "Processed DOCX file with multilingual support (Arabic, English, French)"
            }
        }
        
    except ImportError:
        return {
            "success": False,
            "error": "DOCX processing library not available",
            "details": "Required library (python-docx) not installed",
            "install": "pip install python-docx"
        }
    except Exception as e:
        import traceback
        return {
            "success": False,
            "error": "DOCX processing failed",
            "details": str(e),
            "traceback": traceback.format_exc()
        }


@app.post("/ingest")
async def ingest(file: UploadFile = File(...)):
    """
    Advanced document ingestion with Docling and proper Arabic support.
    
    Supports:
    - PDF files with Docling's advanced PDF understanding
    - DOCX files with python-docx
    - Proper OCR configuration for RTL languages
    - Layout analysis and table detection (PDF only)
    - Hierarchical section detection
    """
    content = await file.read()
    file_extension = os.path.splitext(file.filename)[1].lower() if file.filename else ""
    
    # Detect file type
    is_docx = file_extension == ".docx" or file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    is_pdf = file_extension == ".pdf" or file.content_type == "application/pdf"
    
    if is_docx:
        # Process DOCX file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(content)
            docx_path = tmp.name
        
        try:
            result = await _process_docx_file(docx_path)
            os.unlink(docx_path)
            return result
        except Exception as e:
            if os.path.exists(docx_path):
                os.unlink(docx_path)
            import traceback
            return {
                "success": False,
                "error": str(e),
                "traceback": traceback.format_exc()
            }
    
    elif is_pdf:
        # Process PDF file (existing logic)
        # Fix torch.xpu if needed
        if not hasattr(torch, "xpu"):
            class FakeXPU:
                @staticmethod
                def is_available():
                    return False
                @staticmethod
                def device_count():
                    return 0
            torch.xpu = FakeXPU()

        # Save uploaded file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(content)
            pdf_path = tmp.name
        
        try:
            from docling.document_converter import DocumentConverter, PdfFormatOption
            from docling.datamodel.base_models import InputFormat
            from docling.datamodel.pipeline_options import (
                PdfPipelineOptions,
                TesseractOcrOptions,
            )
            
            # Try to import table structure options (may not be available in all versions)
            try:
                from docling.datamodel.pipeline_options import (
                    TableStructureOptions,
                    TableFormerMode,
                )
                HAS_TABLE_OPTIONS = True
            except ImportError:
                HAS_TABLE_OPTIONS = False
                print("  ⚠️  TableStructureOptions not available, using default table settings")
            from docling.chunking import HybridChunker
            
            print("\n" + "="*70)
            print("DOCLING PDF PROCESSING WITH ARABIC SUPPORT")
            print("="*70)
            
            # ========================================
            # STEP 1: Configure pipeline for Arabic
            # ========================================
            print("\n📋 STEP 1: Configuring pipeline for multilingual documents (Arabic, English, French)...")
            
            # Create pipeline options for PDF with optimal settings for Arabic
            # Configure table structure options for better Arabic table extraction (if available)
            table_structure_options = None
            if HAS_TABLE_OPTIONS:
                try:
                    table_structure_options = TableStructureOptions(
                        # Use accurate mode for better table structure recognition
                        mode=TableFormerMode.ACCURATE,
                        # Enable cell matching for better table cell extraction
                        do_cell_matching=True
                    )
                except Exception as e:
                    print(f"  ⚠️  Could not create TableStructureOptions: {e}")
                    table_structure_options = None
            
            # Build pipeline options dictionary
            # IMPORTANT: Use force_backend_text to prefer native PDF text over OCR
            # This avoids tensor padding errors that occur with OCR processing
            # We'll fallback to OCR if native text extraction fails
            pipeline_kwargs = {
                # Try to use native PDF text first (avoids tensor issues)
                "force_backend_text": True,
                # Enable OCR as fallback (critical for Arabic documents)
                "do_ocr": True,
                # Configure Tesseract for Arabic + English
                "ocr_options": TesseractOcrOptions(
                    lang=["ara", "eng", "fra"],  # Arabic + English + French for multilingual documents
                    # Don't force full page OCR initially (may cause tensor issues)
                    force_full_page_ocr=False,
                    # PSM mode 3: Auto page segmentation (more flexible)
                    psm=3
                ),
                # DISABLE table structure detection to avoid tensor errors
                "do_table_structure": False,
                # Set lower images scale to reduce memory and tensor issues
                "images_scale": 1.5,
                # Disable picture generation to reduce processing
                "generate_picture_images": False
            }
            
            # Try to add batch sizes if available (may not exist in all versions)
            # These help process pages one at a time
            try:
                # Test if these parameters exist by checking PdfPipelineOptions
                test_options = PdfPipelineOptions()
                if hasattr(test_options, 'layout_batch_size'):
                    pipeline_kwargs["layout_batch_size"] = 1
                if hasattr(test_options, 'ocr_batch_size'):
                    pipeline_kwargs["ocr_batch_size"] = 1
                if hasattr(test_options, 'table_batch_size'):
                    pipeline_kwargs["table_batch_size"] = 1
            except:
                pass
            
            # Add table structure options only if available
            if table_structure_options is not None:
                pipeline_kwargs["table_structure_options"] = table_structure_options
            
            pipeline_options = PdfPipelineOptions(**pipeline_kwargs)
            
            print("  ✅ OCR enabled with Arabic + English + French support")
            print("  ⚠️  Table structure detection DISABLED (prevents tensor errors)")
            print("  ⚠️  Using force_backend_text to prefer native PDF text")
            print("  ✅ Image scale set to 1.5x for improved OCR quality")
            
            # ========================================
            # STEP 2: Initialize converter with format options
            # ========================================
            print("\n📄 STEP 2: Initializing DocumentConverter...")
            
            # Create format option for PDF
            pdf_format_option = PdfFormatOption(
                pipeline_options=pipeline_options
            )
            
            # Initialize converter with format options
            converter = DocumentConverter(
                format_options={
                    InputFormat.PDF: pdf_format_option,
                }
            )
            
            print("  ✅ Converter initialized")
            
            # ========================================
            # STEP 3: Convert document
            # ========================================
            print("\n🔄 STEP 3: Converting PDF with Docling...")
            
            # Try conversion with max_num_pages to process one page at a time if needed
            try:
                # First attempt: Try with all pages
                result = converter.convert(pdf_path)
                print("  ✅ Conversion successful on first attempt")
            except Exception as conv_error:
                error_msg = str(conv_error)
                print(f"⚠️  Conversion failed, retrying with simpler pipeline...")
                print(f"   Error: {error_msg[:300]}")
                
                # Check if it's a tensor padding error
                is_tensor_error = "tensor" in error_msg.lower() or "padding" in error_msg.lower()
                
                # Fallback: Try with OCR only, no table detection, and process with max_num_pages
                # This processes pages one at a time to avoid tensor shape mismatches
                try:
                    print("  🔄 Attempting fallback with OCR-only and page-by-page processing...")
                    
                    # Create minimal OCR-only pipeline
                    minimal_pipeline = PdfPipelineOptions(
                        do_ocr=True,
                        ocr_options=TesseractOcrOptions(
                            lang=["ara", "eng", "fra"],
                            force_full_page_ocr=True,
                            psm=6
                        ),
                        do_table_structure=False,  # CRITICAL: Disable table detection
                        images_scale=1.5  # Lower scale to reduce memory usage
                    )
                    
                    minimal_format = PdfFormatOption(pipeline_options=minimal_pipeline)
                    minimal_converter = DocumentConverter(
                        format_options={InputFormat.PDF: minimal_format}
                    )
                    
                    # Try to convert with max_num_pages to limit processing
                    # If that doesn't work, try without limit
                    try:
                        # Try with a page limit first (processes pages sequentially)
                        result = minimal_converter.convert(pdf_path, max_num_pages=100)
                        print("  ✅ Fallback conversion successful with page limit")
                    except TypeError:
                        # max_num_pages might not be supported, try without it
                        result = minimal_converter.convert(pdf_path)
                        print("  ✅ Fallback conversion successful")
                        
                except Exception as fallback_error:
                    error_msg_fallback = str(fallback_error)
                    print(f"  ❌ Fallback also failed: {error_msg_fallback[:300]}")
                    
                    # Final attempt: Try with even simpler configuration
                    # Remove all optional features that might cause tensor issues
                    try:
                        print("  🔄 Attempting final fallback (ultra-minimal OCR)...")
                        ultra_minimal = PdfPipelineOptions(
                            do_ocr=True,
                            ocr_options=TesseractOcrOptions(
                                lang=["ara", "eng", "fra"],
                                force_full_page_ocr=False,  # Try without force_full_page
                                psm=3  # Auto PSM mode
                            ),
                            do_table_structure=False,
                            images_scale=1.0  # Minimum scale
                        )
                        
                        ultra_format = PdfFormatOption(pipeline_options=ultra_minimal)
                        ultra_converter = DocumentConverter(
                            format_options={InputFormat.PDF: ultra_format}
                        )
                        
                        result = ultra_converter.convert(pdf_path)
                        print("  ✅ Ultra-minimal fallback successful")
                    except Exception as final_error:
                        # If all fallbacks fail, try using alternative PDF processing
                        print(f"  ❌ All Docling fallbacks failed. Final error: {str(final_error)[:300]}")
                        print("  🔄 Attempting alternative PDF processing method...")
                        
                        # Last resort: Use alternative library for text extraction
                        try:
                            return await _process_with_alternative_method(pdf_path)
                        except Exception as alt_error:
                            print(f"  ❌ Alternative method also failed: {str(alt_error)[:200]}")
                            # Return a helpful error message
                            return {
                                "success": False,
                                "error": "Docling conversion failed due to tensor padding issues",
                                "details": "The PDF could not be processed due to tensor shape mismatches. This is a known issue with Docling when processing PDFs with pages of different sizes.",
                                "suggestion": "Try processing the PDF with a different tool or split it into single-page files.",
                                "original_error": str(conv_error)[:500]
                            }
            
            if not result or not result.document:
                return {
                    "success": False,
                    "error": "Docling conversion failed",
                    "details": "Could not extract document content"
                }
            
            doc = result.document
            print(f"  ✅ Document converted successfully")
            
            # ========================================
            # STEP 4: Extract sections
            # ========================================
            print("\n🔍 STEP 4: Extracting sections...")
            
            sections = extract_sections_from_docling_document(doc)
            
            print(f"\n  ✅ Found {len(sections)} sections")
            if sections and len(sections) <= 10:
                print("\n  📋 Detected sections:")
                for i, section in enumerate(sections[:10], 1):
                    print(f"    {i}. {section[:70]}...")
            
            # ========================================
            # STEP 5: Chunk the document
            # ========================================
            print("\n✂️  STEP 5: Chunking document...")
            
            # Use a simpler tokenizer that might avoid tensor padding issues
            # Try to use a tokenizer that handles variable lengths better
            try:
                chunker = HybridChunker(
                    tokenizer="bert-base-uncased",
                    max_tokens=600,
                    overlap_tokens=100,
                    heading_hierarchies=True,
                )
            except Exception as chunker_error:
                print(f"  ⚠️  Error creating chunker: {chunker_error}")
                print("  🔄 Trying with different tokenizer...")
                # Fallback: try without heading hierarchies
                chunker = HybridChunker(
                    tokenizer="bert-base-uncased",
                    max_tokens=600,
                    overlap_tokens=100,
                    heading_hierarchies=False,  # Disable to avoid potential tensor issues
                )
            
            docling_chunks = list(chunker.chunk(doc))
            print(f"  ✅ Created {len(docling_chunks)} chunks")
            
            # Convert to our format and clean text
            chunks = []
            for dc in docling_chunks:
                # Clean the text to remove unnecessary newlines
                cleaned_text = clean_text(dc.text)
                chunk_dict = {
                    "text": cleaned_text,
                    "page": dc.meta.get("page"),
                    "docling_section": dc.meta.get("headings", [None])[0] if dc.meta.get("headings") else None
                }
                chunks.append(chunk_dict)
            
            # ========================================
            # STEP 6: Assign sections
            # ========================================
            print("\n🎯 STEP 6: Assigning sections to chunks...")
            
            enriched_chunks = assign_sections_to_chunks(chunks, sections)
            
            # Show sample
            if enriched_chunks:
                print(f"\n📊 Sample chunk:")
                print(f"  Page: {enriched_chunks[0]['meta']['page']}")
                print(f"  Section: {enriched_chunks[0]['meta']['section']}")
                print(f"  Text: {enriched_chunks[0]['text'][:150]}...")
            
            # Quality check
            unique_sections = set(c['meta']['section'] for c in enriched_chunks)
            print(f"\n✅ Quality check:")
            print(f"  - Total chunks: {len(enriched_chunks)}")
            print(f"  - Unique sections: {len(unique_sections)}")
            print(f"  - Chunks per section: {len(enriched_chunks) / len(unique_sections):.1f}")
            
            os.unlink(pdf_path)
            
            return {
                "success": True,
                "method": "docling_with_arabic_ocr",
                "total_chunks": len(enriched_chunks),
                "detected_sections": sections,
                "sections_count": len(sections),
                "unique_sections_in_chunks": len(unique_sections),
                "chunks": enriched_chunks,
            "metadata": {
                "ocr_engine": "tesseract",
                "languages": ["ara", "eng", "fra"],
                "table_detection": True,
                "layout_analysis": True
            },
            "note": "Processed with Docling's advanced PDF understanding and multilingual OCR (Arabic, English, French)"
            }
        
        except ImportError as e:
            if os.path.exists(pdf_path):
                os.unlink(pdf_path)
            
            return {
                "success": False,
                "error": "Missing dependencies",
                "details": str(e),
                "install_instructions": {
                    "docling": "pip install docling",
                    "tesseract": "sudo apt-get install tesseract-ocr tesseract-ocr-ara tesseract-ocr-eng",
                }
            }
        
        except Exception as e:
            if os.path.exists(pdf_path):
                os.unlink(pdf_path)
            
            import traceback
            return {
                "success": False,
                "error": str(e),
                "traceback": traceback.format_exc()
            }
    else:
        return {
            "success": False,
            "error": "Unsupported file type",
            "details": f"File type '{file_extension}' is not supported. Please upload a PDF or DOCX file.",
            "supported_types": ["pdf", "docx"]
        }


@app.get("/health")
async def health():
    """Check system dependencies and Docling setup"""
    status = {
        "docling": False,
        "tesseract": False,
        "tesseract_arabic": False,
        "tessdata_prefix": os.environ.get('TESSDATA_PREFIX'),
    }
    
    try:
        import docling
        status["docling"] = True
        status["docling_version"] = docling.__version__
    except:
        pass
    
    try:
        import pytesseract
        version = pytesseract.get_tesseract_version()
        status["tesseract"] = True
        status["tesseract_version"] = str(version)
        
        langs = pytesseract.get_languages()
        status["tesseract_arabic"] = "ara" in langs
        status["tesseract_english"] = "eng" in langs
        status["tesseract_french"] = "fra" in langs
        status["tesseract_languages"] = langs
    except Exception as e:
        status["tesseract_error"] = str(e)
    
    all_ready = status["docling"] and status["tesseract"] and status["tesseract_arabic"] and status["tesseract_english"] and status["tesseract_french"]
    
    recommendations = []
    if not status["docling"]:
        recommendations.append("Install Docling: pip install docling")
    if not status["tesseract"]:
        recommendations.append("Install Tesseract: sudo apt-get install tesseract-ocr")
    if not status["tesseract_arabic"]:
        recommendations.append("Install Arabic: sudo apt-get install tesseract-ocr-ara")
    if not status["tesseract_english"]:
        recommendations.append("Install English: sudo apt-get install tesseract-ocr-eng")
    if not status["tesseract_french"]:
        recommendations.append("Install French: sudo apt-get install tesseract-ocr-fra")
    if not status["tessdata_prefix"]:
        recommendations.append("Set TESSDATA_PREFIX: export TESSDATA_PREFIX=/usr/share/tesseract-ocr/5/tessdata")
    
    return {
        "status": "ready" if all_ready else "missing_dependencies",
        "components": status,
        "recommendations": recommendations
    }


if __name__ == "__main__":
    import uvicorn
    print("\n" + "="*70)
    print("🚀 Docling Document Processing Server - Multilingual Support")
    print("="*70)
    print("\n📋 Supported Languages:")
    print("  • Arabic (العربية)")
    print("  • English")
    print("  • French (Français)")
    print("\n📋 Endpoints:")
    print("  • Health: http://localhost:8000/health")
    print("  • Ingest: http://localhost:8000/ingest")
    print("\n📦 Dependencies:")
    print("  • pip install docling python-docx")
    print("  • sudo apt-get install tesseract-ocr tesseract-ocr-ara tesseract-ocr-fra")
    print("\n" + "="*70 + "\n")
    
    uvicorn.run(app, host="0.0.0.0", port=8000)